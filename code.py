from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import json


def capture_formatting(paragraph):
    """Capture formatting from a paragraph's first run."""
    if not paragraph.runs:
        return None

    first_run = paragraph.runs[0]
    return {
        'bold': first_run.bold,
        'italic': first_run.italic,
        'underline': first_run.underline,
        'font_size': first_run.font.size,
        'font_name': first_run.font.name,
        'font_color': first_run.font.color.rgb if first_run.font.color.rgb else None
    }


def apply_formatting(run, formatting):
    """Apply formatting to a run."""
    if not formatting:
        return

    if formatting['bold'] is not None:
        run.bold = formatting['bold']
    if formatting['italic'] is not None:
        run.italic = formatting['italic']
    if formatting['underline'] is not None:
        run.underline = formatting['underline']
    if formatting['font_size']:
        run.font.size = formatting['font_size']
    if formatting['font_name']:
        run.font.name = formatting['font_name']
    if formatting['font_color']:
        run.font.color.rgb = formatting['font_color']


def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Replace text in a paragraph while preserving formatting."""
    if old_text not in paragraph.text:
        return False

    # Try to replace within a single run
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True

    # If placeholder spans multiple runs, rebuild paragraph
    full_text = paragraph.text
    if old_text in full_text:
        first_run = paragraph.runs[0] if paragraph.runs else None

        # Clear all runs
        for run in paragraph.runs:
            run.text = ''

        # Create new run with replaced text
        new_run = paragraph.add_run(full_text.replace(old_text, new_text))

        # Copy formatting from first run
        if first_run:
            new_run.bold = first_run.bold
            new_run.italic = first_run.italic
            new_run.underline = first_run.underline
            new_run.font.size = first_run.font.size
            new_run.font.name = first_run.font.name
            if first_run.font.color.rgb:
                new_run.font.color.rgb = first_run.font.color.rgb

        return True

    return False


def create_bullet_paragraph(doc, paragraph_index, text, formatting=None):
    """Create a single bullet paragraph with proper formatting."""
    parent = doc.paragraphs[paragraph_index]._element.getparent()

    # Create new paragraph
    new_para = doc.add_paragraph()
    run = new_para.add_run(text)

    # Apply text formatting
    apply_formatting(run, formatting)

    # Set indentation
    new_para.paragraph_format.left_indent = Inches(0.13)
    new_para.paragraph_format.right_indent = Inches(0)
    new_para.paragraph_format.first_line_indent = Inches(-0.13)

    # Set spacing
    new_para.paragraph_format.space_before = Pt(0)
    new_para.paragraph_format.space_after = Pt(0)
    new_para.paragraph_format.line_spacing = 1.0

    # Add bullet formatting using XML
    pPr = new_para._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')

    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numPr.append(ilvl)

    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    numPr.append(numId)

    pPr.append(numPr)

    # Insert at correct position
    parent.insert(paragraph_index + 1, new_para._element)


def insert_bullets(doc, paragraph_index, bullets, formatting=None):
    """Insert multiple bullet points after a paragraph."""
    for i, bullet in enumerate(bullets):
        create_bullet_paragraph(doc, paragraph_index + i, bullet, formatting)


def process_simple_replacements(doc, json_data):
    """Process all simple text replacements (non-bullet fields)."""
    # Define all simple replacement mappings
    replacements = {}
    
    # Experience section
    if 'experience' in json_data:
        exp = json_data['experience']
        replacements['[experience_1_title_company_location]'] = exp.get('experience_1_title_company_location', '')
        replacements['[experience_1_duration]'] = exp.get('experience_1_duration', '')
        replacements['[experience_2_title_company_location]'] = exp.get('experience_2_title_company_location', '')
        replacements['[experience_2_duration]'] = exp.get('experience_2_duration', '')
    
    # Projects section
    if 'projects' in json_data:
        proj = json_data['projects']
        replacements['[project_1_title]'] = proj.get('project_1_title', '')
        replacements['[project_1_duration]'] = proj.get('project_1_duration', '')
        replacements['[project_2_title]'] = proj.get('project_2_title', '')
        replacements['[project_2_duration]'] = proj.get('project_2_duration', '')
    
    # Education section
    if 'education' in json_data:
        edu = json_data['education']
        replacements['[masters_coursework]'] = edu.get('masters_coursework', '')
    
    # Summary section
    if 'summary' in json_data:
        summary_data = json_data['summary']
        replacements['[summary_text]'] = summary_data.get('summary_text', '')
    
    # Perform all simple replacements
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            replace_text_in_paragraph(paragraph, placeholder, value)


def process_bullet_sections(doc, json_data):
    """Process sections with bullet points."""
    paragraphs_to_process = []
    
    # Collect all paragraphs with bullet placeholders
    for i, paragraph in enumerate(doc.paragraphs):
        if '[experience_1_bullets]' in paragraph.text:
            paragraphs_to_process.append({
                'index': i,
                'placeholder': '[experience_1_bullets]',
                'data': json_data.get('experience', {}).get('experience_1_bullets', [])
            })
        elif '[experience_2_bullets]' in paragraph.text:
            paragraphs_to_process.append({
                'index': i,
                'placeholder': '[experience_2_bullets]',
                'data': json_data.get('experience', {}).get('experience_2_bullets', [])
            })
        elif '[project_1_bullets]' in paragraph.text:
            paragraphs_to_process.append({
                'index': i,
                'placeholder': '[project_1_bullets]',
                'data': json_data.get('projects', {}).get('project_1_bullets', [])
            })
        elif '[project_2_bullets]' in paragraph.text:
            paragraphs_to_process.append({
                'index': i,
                'placeholder': '[project_2_bullets]',
                'data': json_data.get('projects', {}).get('project_2_bullets', [])
            })
        elif '[achievements_bullets]' in paragraph.text:
            paragraphs_to_process.append({
                'index': i,
                'placeholder': '[achievements_bullets]',
                'data': json_data.get('achievements', {}).get('achievements_bullets', [])
            })
    
    # Process in reverse order to maintain correct indices
    for item in reversed(paragraphs_to_process):
        paragraph = doc.paragraphs[item['index']]
        formatting = capture_formatting(paragraph)
        
        # Remove the placeholder paragraph
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
        
        # Insert bullets
        insert_bullets(doc, item['index'] - 1, item['data'], formatting)


def process_skills(doc, skills_data):
    """Process skills section."""
    for i, paragraph in enumerate(doc.paragraphs):
        if '[skills_category_list]' not in paragraph.text:
            continue

        # Capture formatting
        ref_font_name = None
        ref_font_size = None
        ref_font_color = None
        style = paragraph.style

        if paragraph.runs:
            ref_font = paragraph.runs[0].font
            ref_font_name = ref_font.name
            ref_font_size = ref_font.size
            ref_font_color = ref_font.color.rgb if ref_font.color.rgb else None

        # Clear placeholder
        paragraph.text = ""
        p_index = i

        # Insert skills
        for group in skills_data:
            new_p = doc.paragraphs[p_index].insert_paragraph_before()
            new_p.style = style
            new_p.paragraph_format.space_after = 0

            # Bold category
            run_bold = new_p.add_run(f"{group['category']}: ")
            run_bold.bold = True
            if ref_font_name:
                run_bold.font.name = ref_font_name
            if ref_font_size:
                run_bold.font.size = ref_font_size
            if ref_font_color:
                run_bold.font.color.rgb = ref_font_color

            # Regular items
            run_items = new_p.add_run(", ".join(group["items"]))
            if ref_font_name:
                run_items.font.name = ref_font_name
            if ref_font_size:
                run_items.font.size = ref_font_size
            if ref_font_color:
                run_items.font.color.rgb = ref_font_color

            p_index += 1

        # Remove placeholder paragraph
        paragraph._element.getparent().remove(paragraph._element)
        break


def generate_resume(template_path, output_path, json_input_path):
    """Main function to generate resume from template and JSON data."""
    # Load JSON data from file
    with open(json_input_path, 'r', encoding='utf-8') as f:
        json_data = json.load(f)
    
    # Load document
    doc = Document(template_path)
    
    # Process simple text replacements first
    process_simple_replacements(doc, json_data)
    
    # Process bullet sections
    process_bullet_sections(doc, json_data)
    
    # Process skills section
    if 'skills' in json_data and 'skills_category_list' in json_data['skills']:
        process_skills(doc, json_data['skills']['skills_category_list'])
    
    # Save document
    doc.save(output_path)
    print(f"✅ Resume generated successfully: {output_path}")


# Example usage
if __name__ == "__main__":
    template_file = 'template.docx'
    output_file = 'resume_filled.docx'
    json_input_file = 'json_input.json'
    
    generate_resume(template_file, output_file, json_input_file)